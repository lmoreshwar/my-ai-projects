from docx import Document
from typing import Dict, Any
import os

class DocxGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template not found at {self.template_path}")

    def generate(self, data: Dict[str, Any], output_path: str):
        """Fill the docx template with provided data."""
        doc = Document(self.template_path)
        
        # Mapping logic:
        # We look for placeholders like {{PROJECT_NAME}}, {{SUMMARY}}, etc.
        # Or we just append content to the end for simplicity in this version.
        
        # Simple replacement in paragraphs
        for p in doc.paragraphs:
            for key, val in data.items():
                placeholder = "{{" + key.upper() + "}}"
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(val))
        
        # Simple replacement in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in data.items():
                        placeholder = "{{" + key.upper() + "}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(val))
        
        # If the template is empty or needs a new section:
        if "test_plan_content" in data:
            doc.add_heading("Generated Test Plan Content", level=1)
            doc.add_paragraph(data["test_plan_content"])

        doc.save(output_path)
        return output_path
