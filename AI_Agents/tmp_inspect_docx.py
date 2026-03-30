import os
from docx import Document

template_path = r"c:\Users\DELL\AI Workspace\AI My Projects\MyAIProjects\AI_Agents\test_plan_document\Test Plan - Template.docx"

if os.path.exists(template_path):
    try:
        doc = Document(template_path)
        print("--- Paragraphs in Template ---")
        for i, p in enumerate(doc.paragraphs):
            if "{{" in p.text or "<<" in p.text or "[" in p.text:
                print(f"L{i}: {p.text}")
        
        print("\n--- Tables in Template ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
            # Print first row of each table to identify it
            if len(table.rows) > 0:
                print(f"  Header: {[cell.text.strip() for cell in table.rows[0].cells]}")
    except Exception as e:
        print(f"Error reading docx: {e}")
else:
    print("Template not found.")
